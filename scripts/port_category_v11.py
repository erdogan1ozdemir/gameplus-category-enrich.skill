# -*- coding: utf-8 -*-
"""28 GFN kategori içeriğini v11'e taşır ve her biri için .docx üretir.

Kural: YAZARIN METNİNE DOKUNULMAZ. Görsel katman yenilenir, bölüm sırası
kapanış -> Fix CTA -> SSS olur, FAQ Schema eklenir, gövdeye 1-2 kategori linki konur.

Kullanım: python3 kategori_batch_v11.py [slug ...]   (argümansız: hepsi)
"""
import os
import re
import json
import sys

SKILLDIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SKILLDIR)
from category_components import *  # noqa: F401,F403
sys.path.insert(0, os.path.join(SKILLDIR, "..", "..", "gameplus-blog-enrich-v2", "scripts"))
from gameplus_blog_components import apply_link_policy   # Kural 21
from category_components import render_kategori_schema
from gameplus_blog_components import auto_link_categories
from icerik_gelistirme import ALT_TUR_TANIM, EK_FAQ, KATEGORI_TANIM
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor

KOK = os.environ.get("GAMEPLUS_OUT", os.getcwd())
KAYNAK = os.path.join(KOK, "kategori-kaynak")
CIKTI = os.path.join(KOK, "kategori-icerikleri")

# Sayfa-ortasi CTA dagilimi (kullanici onayi): agirlikli /paketler, magazalar /geforce-now-nedir.
PAKETLER = "https://gameplus.com.tr/paketler"
NEDIR = "https://gameplus.com.tr/geforce-now-nedir"
CTA_HEDEF = {s: NEDIR for s in
             ("steam", "xbox", "epic-games", "ea-app", "ubisoft-connect", "gog", "demo")}


# Kaynakta lisans hatirlatmasi yoksa standart metin eklenir (bu bizim enrichment
# bilesenimiz, yazarin metni degil; skill zorunlu tutuyor).
LISANS_METNI = ("GeForce NOW oyun satmaz; mevcut Steam, Epic Games Store, Xbox, EA App veya "
                "Ubisoft Connect kütüphanendeki lisansları bulut üzerinden çalıştırır. Bir yapımı "
                "oynamak için ilgili mağazada geçerli bir lisansa ya da aboneliğe sahip olman gerekir.")

# Kaynakta sayfa-ortasi CTA yoksa kategoriye uygun metinle eklenir.
DYN_CTA_METNI = {
    "bulmaca": ("Bulmaca çözmek için güçlü bir bilgisayar gerekmiyor.",
                "GAME+ paketleriyle kütüphanendeki bulmaca ve zeka oyunlarını uyumlu cihazında, "
                "kurulum beklemeden açabilirsin."),
    "strateji": ("Strateji oyunları uzun oturum ister, güçlü donanım değil.",
                 "GAME+ paketleriyle kütüphanendeki strateji yapımlarını bulut üzerinden "
                 "çalıştırabilirsin."),
    "epic-games": ("Epic Games kütüphanen bulutta seni bekliyor.",
                   "Hesabını bağladıktan sonra desteklenen yapımları indirme beklemeden "
                   "açabilirsin; bulut oyunun nasıl çalıştığını inceleyebilirsin."),
    "bagimsiz": ("İndie yapımlar için güçlü donanım şart değil.",
                 "GAME+ paketleriyle kütüphanendeki bağımsız oyunları uyumlu cihazında, kurulum "
                 "beklemeden açabilirsin."),
    "basit-eglence": ("Kısa molalarda açıp kapatabileceğin oyunlar.",
                      "GAME+ paketleriyle casual yapımları indirme beklemeden başlatabilirsin."),
    "moba": ("5v5 maçlarda gecikme belirleyici olur.",
             "GAME+ paketleriyle MOBA oyunlarını Türkiye sunucuları üzerinden oynayabilirsin."),
    "oyunlar": ("Kütüphanen bulutta, kurulum yok.",
                "GAME+ paketleriyle desteklenen oyunlarını uyumlu cihazında açabilirsin."),
    "yaris": ("Direksiyon başına geçmek için güçlü bir PC gerekmiyor.",
              "GAME+ paketleriyle kütüphanendeki yarış oyunlarını bulut üzerinden çalıştırabilirsin."),
    "gog": ("GOG kütüphaneni buluta taşı.",
            "Hesap bağlama ve desteklenen oyunların bulutta nasıl çalıştığı hakkında ayrıntıları "
            "inceleyebilirsin."),
}
_VARSAYILAN_DYN = ("Bulutta oyna, indirme bekleme.",
                   "GAME+ paketleriyle kütüphanendeki oyunları uyumlu cihazında açabilirsin.")



# FAQ sorularina kategori baglami eklenir (kullanici revizesi): soru tek basina okundugunda
# hangi kategoriye ait oldugu anlasilmali, yoksa genel bir soruya donuyor ve hedef query'ye
# oturmuyor. YALNIZ baglamsiz kalanlar degistirildi; oyun adi geceni ya da kategorinin kendi
# terimini tasiyanlar (RPG, casual, indie, metroidvania...) aynen birakildi.
FAQ_SORU_REVIZE = {
    ("aile-dostu", "GeForce NOW'da çocuklarla oynanabilecek oyunlar var mı?"):
        "GeForce NOW'da aile dostu, çocuklarla oynanabilecek oyunlar var mı?",
    ("aksiyon", "Açık dünya oyunları bulutta kasma yapar mı?"):
        "Açık dünya aksiyon oyunları bulutta kasma yapar mı?",
    ("arcade", "Emülatör ve ROM dosyaları GeForce NOW'da çalışır mı?"):
        "Arcade oyunları için emülatör ve ROM dosyaları GeForce NOW'da çalışır mı?",
    ("bulmaca", "Zeka geliştiren oyunlar GeForce NOW'da var mı?"):
        "Zeka geliştiren bulmaca oyunları GeForce NOW'da var mı?",
    ("dovus-oyunu", "Fight stick GeForce NOW'da çalışır mı?"):
        "Dövüş oyunlarında fight stick GeForce NOW'da çalışır mı?",
    ("fps", "Klavye-fare mı, gamepad mi tercih edilmeli?"):
        "FPS oyunlarında klavye-fare mı, gamepad mi tercih edilmeli?",
    ("macera", "Açık dünya oyunları nedir?"):
        "Açık dünya macera oyunları nedir?",
    ("macera", "GTA tarzı açık dünya oyunları nelerdir?"):
        "GTA tarzı açık dünya macera oyunları nelerdir?",
    ("macera", "Hikaye odaklı en iyi oyunlar nelerdir?"):
        "Hikaye odaklı en iyi macera oyunları nelerdir?",
    ("oynamasi-ucretsiz", "Oyun içi satın alımlar ve ilerleme senkronize olur mu?"):
        "Ücretsiz oyunlarda oyun içi satın alımlar ve ilerleme senkronize olur mu?",
    ("simulasyon", "Direksiyon seti ve HOTAS joystick GeForce NOW'da çalışır mı?"):
        "Simülasyon oyunlarında direksiyon seti ve HOTAS joystick GeForce NOW'da çalışır mı?",
    ("xbox", "Microsoft hesabımı nasıl bağlarım?"):
        "Xbox oyunları için Microsoft hesabımı GeForce NOW'a nasıl bağlarım?",
    ("yaris", "Direksiyon seti ile GeForce NOW'da oynanabilir mi?"):
        "Yarış oyunları direksiyon seti ile GeForce NOW'da oynanabilir mi?",
    ("yaris", "Dünyanın en gerçekçi araba oyunu nedir?"):
        "Dünyanın en gerçekçi araba yarışı oyunu nedir?",
}

# Tazelik sinyali: uretim tarihi. Uydurma tarih verilmez; icerik yenilendiginde guncellenir.
URETIM_TARIHI = "2026-09-02"

# Kullanici onayli baslik dilbilgisi duzeltmeleri (yazarin basligi; "-ini nasil ...nir" edilgen
# kalibinda belirtme eki fazlaydi). Anahtar: (slug, duz baslik metni).
BASLIK_REVIZE = {
    ("epic-games", "GeForce NOW'da Epic Games Hesabını Nasıl Bağlanır?"):
        "GeForce NOW'a Epic Games Hesabı Nasıl Bağlanır?",
    ("gog", "GeForce NOW'da GOG Oyunlarını Nasıl Oynanır?"):
        "GeForce NOW'da GOG Oyunları Nasıl Oynanır?",
    ("steam", "GeForce NOW'da Steam Oyunlarını Nasıl Oynanır?"):
        "GeForce NOW'da Steam Oyunları Nasıl Oynanır?",
}


def yazim_normalize(html):
    """Kullanici karari: 'Indie' Ingilizce sozcuk, Turkce buyuk I ile yazilmaz; 'Arcade' ve
    'Indie' tur adlari buyuk harfle baslar. Yalniz metin dugumlerinde calisir (URL/slug/etiket
    icine dokunmaz). Donus: (html, sayac)."""
    sayac = {"İndie>Indie": 0, "arcade>Arcade": 0, "indie>Indie": 0}
    def _metin(m):
        t = m.group(0)
        n = t.count("İndie"); t = t.replace("İndie", "Indie"); sayac["İndie>Indie"] += n
        t, k = re.subn(r"(?<![\w/.\-])arcade(?!\w)", "Arcade", t); sayac["arcade>Arcade"] += k
        t, k = re.subn(r"(?<![\w/.\-])indie(?!\w)", "Indie", t); sayac["indie>Indie"] += k
        return t
    # etiketler ve script/style blogu disindaki metin: ">...<" araliklari
    parcalar = re.split(r"(<script.*?</script>|<style.*?</style>|<[^>]+>)", html, flags=re.S)
    for i in range(0, len(parcalar), 2):
        parcalar[i] = _metin(re.match(r".*", parcalar[i], re.S))
    # schema JSON icindeki metin alanlari: JSON olarak cozulur, URL disi string degerler duzeltilir
    def _json_yuru(o):
        if isinstance(o, dict):
            return {k: _json_yuru(x) for k, x in o.items()}
        if isinstance(o, list):
            return [_json_yuru(x) for x in o]
        if isinstance(o, str) and not o.startswith("http"):
            return _metin(re.match(r".*", o, re.S))
        return o
    for i in range(1, len(parcalar), 2):
        if parcalar[i].startswith("<script") and "ld+json" in parcalar[i][:80]:
            m = re.match(r"(<script[^>]*>)(.*)(</script>)", parcalar[i], re.S)
            try:
                veri = json.loads(m.group(2))
            except ValueError:
                continue
            parcalar[i] = m.group(1) + "\n" + json.dumps(_json_yuru(veri), ensure_ascii=False, indent=1) + "\n" + m.group(3)
    return "".join(parcalar), sayac

BASE = "https://gameplus.com.tr/gfn/oyunlar"


def kategori_url(slug):
    return BASE if slug == "oyunlar" else f"{BASE}/{slug}"


def ic_html(el):
    kopya = BeautifulSoup(str(el), "html.parser")
    for t in kopya.find_all(True):
        for oz in ("style", "class"):
            t.attrs.pop(oz, None)
    kok = kopya.find(el.name)
    return "".join(str(c) for c in kok.children).strip()


def duz(el):
    return re.sub(r"\s+", " ", el.get_text()).strip()



_ETIKET_ONEK = re.compile(
    r"^\s*(?:ℹ|📝|⚠|💡)?\s*(?:GAME\+|Game\+)?\s*(?:Editör Notu|Editor Notu|Hatırlatma|Not)\s*:?\s*",
    re.I)


def onek_temizle(t):
    """Kaynakta 'ℹ Hatırlatma ...' / '📝 Game+ Editör Notu ...' seklinde etiket metne gomulu
    geliyor; bilesen bu etiketi zaten eyebrow olarak basiyor, tekrari kaldirilir."""
    return _ETIKET_ONEK.sub("", t).lstrip()


def cikar(soup):
    v = {"tldr": [], "info": [], "tablolar": [], "notlar": [], "callout": None,
         "dyn": None, "fix": None, "faq": [], "akis": []}
    for el in soup.children:
        if getattr(el, "name", None) is None or el.name == "style":
            continue
        cls = " ".join(el.get("class", []))
        if el.name in ("h2", "h3", "h4"):
            v["akis"].append((el.name, ic_html(el)))
        elif el.name == "p":
            v["akis"].append(("p", ic_html(el)))
        elif el.name in ("ul", "ol") and "tldr" not in cls:
            v["akis"].append((el.name, [ic_html(li) for li in el.find_all("li", recursive=False)]))
        elif "tldr-block" in cls:
            for li in el.find_all("li"):
                v["tldr"].append(ic_html(li))
            v["akis"].append(("tldr", None))
        elif "info-card" in cls:
            for h in el.find_all("div", recursive=False):
                ogeler = [x for x in h.find_all(["span", "div", "p"]) if duz(x)]
                gor, sec = [], set()
                for x in ogeler:
                    t = duz(x)
                    if t in sec:
                        continue
                    sec.add(t); gor.append(x)
                if len(gor) < 2:
                    continue
                # v9: ETIKET ustte (uppercase span) | v10.2: DEGER ustte (24px sari)
                st0 = (gor[0].get("style") or "").lower()
                deger_ust = ("ffc900" in st0 or "font-size:24px" in st0.replace(" ", "")
                             or "1.7em" in st0)
                etiket_ust = "text-transform:uppercase" in st0.replace(" ", "")
                if deger_ust and not etiket_ust:
                    v["info"].append((duz(gor[1]), duz(gor[0])))   # (etiket, deger)
                else:
                    v["info"].append((duz(gor[0]), duz(gor[1])))
            v["akis"].append(("info", None))
        elif "table-wrap" in cls or el.name == "table" or el.find("table"):
            t = el if el.name == "table" else el.find("table")   # sarmalayicisiz <table> da alinir
            v["tablolar"].append(([ic_html(th) for th in t.find_all("th")],
                                  [[ic_html(td) for td in tr.find_all("td")]
                                   for tr in t.find_all("tr") if tr.find("td")]))
            v["akis"].append(("tablo", len(v["tablolar"]) - 1))
        elif "editor-note" in cls:
            p = el.find("p")
            v["notlar"].append(onek_temizle(ic_html(p) if p else duz(el)))
            v["akis"].append(("not", len(v["notlar"]) - 1))
        elif "callout" in cls or "highlight-box" in cls:
            p = el.find("p")
            v["callout"] = onek_temizle(ic_html(p) if p else duz(el))
            v["akis"].append(("callout", None))
        elif el.name == "details":
            # v10.2 kaynaginda <details> ogeleri sarmalayicisiz, ust duzeyde duruyor
            cevap = el.find("p") or el.find("div")
            v["faq"].append((re.sub(r"^\s*\+\s*", "", duz(el.find("summary"))),
                             ic_html(cevap) if cevap else ""))
            if not any(t == "faq" for t, _ in v["akis"]):
                v["akis"].append(("faq", None))
        elif "faq-block" in cls or el.find("details"):
            for d in el.find_all("details"):
                cevap = d.find("p") or d.find("div")
                v["faq"].append((re.sub(r"^\s*\+\s*", "", duz(d.find("summary"))),
                                 ic_html(cevap) if cevap else ""))
            v["akis"].append(("faq", None))
        else:
            a = el.find("a", href=True)
            if not a:
                continue
            # Yaprak div'leri sirayla topla; istatistik seridi (sayi + etiket ciftleri) ve
            # PERFORMANCE/ULTIMATE rozetleri elenir. Ilk kalan = baslik, ikincisi = aciklama.
            adaylar = []
            for d in el.find_all("div"):
                if d.find("div") or d.find("a"):
                    continue
                t = duz(d)
                if not t or len(t) < 3:
                    continue
                st = (d.get("style") or "").replace(" ", "").lower()
                onceki = d.find_previous_sibling("div")
                st_onceki = ((onceki.get("style") or "").replace(" ", "").lower()
                             if onceki is not None else "")
                # C4 "stat banner" seridi: buyuk puntolu sayi + uppercase etiket ciftleri
                # NOT: "font-weight:800" olcutu kullanilmaz; v10.2 kaynaginda CTA basligi da 800.
                if "font-size:1.7em" in st:
                    continue                                   # istatistik degeri
                if "text-transform:uppercase" in st:
                    continue                                   # istatistik etiketi / rozet
                if "font-size:1.7em" in st_onceki:
                    continue                                   # istatistik etiketi (deger altinda)
                if re.fullmatch(r"[\d.,+%\s]+", t):
                    continue
                if t.isupper():
                    continue                                   # PERFORMANCE / ULTIMATE rozet satiri
                if d.find("span") and not d.find(string=lambda x: isinstance(x, str) and x.strip()
                                                  and x.parent is d):
                    continue                                   # yalniz rozet span'lari iceren div
                if not re.search(r"[.!?:,]", t) and len(adaylar) == 0 and len(t.split()) <= 2:
                    continue                                   # 1-2 kelimelik etiket, baslik degil
                adaylar.append(t)
            p = el.find("p")
            baslik = adaylar[0] if adaylar else ""
            if p is not None:
                aciklama = ic_html(p)
            else:
                aciklama = adaylar[1] if len(adaylar) > 1 else ""
            if not baslik:
                baslik, aciklama = None, None      # kullanilabilir kopya yok -> yedek metin
            # Eski icerikte HER IKI CTA da /gfn/paketler'e gidiyordu (v9 davranisi), bu yuzden
            # href'e gore ayirmak yanlis. Siraya gore ayiriyoruz: SON CTA karti Fix CTA'dir,
            # ondan onceki(ler) sayfa-ortasi dynamic CTA'dir.
            v.setdefault("ctalar", []).append((baslik, aciklama))
            v["akis"].append(("cta", len(v["ctalar"]) - 1))

    ctalar = v.get("ctalar", [])
    if ctalar:
        v["fix"] = ctalar[-1]
        ilk = ctalar[0] if len(ctalar) > 1 else None
        v["dyn"] = ilk if (ilk and ilk[0]) else None
        son = len(ctalar) - 1
        v["akis"] = [(("fix" if d == son else ("dyn" if d == 0 else "cta_atla")), None)
                     if t == "cta" else (t, d) for t, d in v["akis"]]
    return v



# =============================================================================
# MOD IDDIALARI KALDIRILDI (kullanici karari)
# Bulut oturumunda dosya sistemine erisim yok ve mod destegi oyundan oyuna
# degisiyor; ozellikle "sonraki oturumlarda korunur" bir kaliCILIK VAADI.
# NOT: Turkcede "mod" hem game mode hem modification demek. Burada YALNIZ
# modification iddialari kaldirilir; "rekabetci modlar", "cok oyunculu modlar"
# gibi OYUN MODU cumlelerine DOKUNULMAZ.
# =============================================================================

MOD_CUMLE = {
    "steam": [
        ("Achievement, Workshop ve oyun süresi takip eder.",
         "Achievement ve oyun süresi takip eder."),
        ("Modlar oturum içinde yüklenir, sonraki oturumlarda korunur.", ""),
        ("Cities: Skylines II, Civilization VII, Crusader Kings III, Farming Simulator 25 ve "
         "benzeri Workshop destekli Steam oyunlarında topluluk modları doğrudan çalışır.", ""),
        ("Modlar oturum içinde Steam üzerinden yüklenir ve sonraki oturumlarda korunur.", ""),
        ("Yerel dosya sistemiyle doğrudan etkileşim gerektiren bazı modlar bulut ortamında "
         "sınırlı olabilir.", ""),
    ],
    "canlandirma": [
        ("Skyrim ve Pathfinder gibi mod ağırlıklı yapımlarda topluluk modları bulutta çalışır.", ""),
    ],
    "simulasyon": [
        ("Cities: Skylines II, Farming Simulator 25 ve Anno 117 gibi yapımlarda topluluk modları "
         "aktif olarak çalışır.", ""),
    ],
    "bagimsiz": [
        ("Mod desteği bulutta.", ""),
        ("Mod kuruluşu oturum içinde gerçekleşir ve sonraki oturumlarda korunur.", ""),
    ],
    "strateji": [
        ("Crusader Kings III ve Civilization serisi gibi oyunlarda topluluk modları aktif olarak "
         "çalışır.", ""),
        ("Modlar oturum içinde yüklenir ve sonraki oturumlarda da korunur.", ""),
    ],
}

# Tamami mod iddiasi olan SSS maddeleri: soru + cevap birlikte dusurulur,
# boylece FAQ Schema'dan da cikar.
FAQ_KALDIR = {
    "strateji": ["Strateji oyunlarına GeForce NOW'da mod yüklenebilir mi?"],
}


def mod_iddialarini_kaldir(html, slug):
    """Govdeden modification iddialarini cikarir. Bosalan <li>/<p> ogeleri de atilir."""
    degisim = []
    for eski, yeni in MOD_CUMLE.get(slug, []):
        if eski in html:
            html = html.replace(eski, yeni)
            degisim.append((eski[:60] + ("…" if len(eski) > 60 else ""), yeni or "(kaldırıldı)"))
    # bosalan ogeler
    html = re.sub(r"<li>\s*(?:<span[^>]*></span>)?\s*<span>\s*</span>\s*</li>\s*", "", html)
    html = re.sub(r"<li>\s*</li>\s*", "", html)
    html = re.sub(r"<p>\s*</p>\s*", "", html)
    html = re.sub(r"(<td[^>]*>)\s+", r"\1", html)
    html = re.sub(r"[ \t]{2,}", " ", html)
    return html, degisim


def kur(v, slug):
    dyn_url = CTA_HEDEF.get(slug, PAKETLER)
    parcalar, kaynak = [], []

    i_sss = next((i for i, (t, _) in enumerate(v["akis"]) if t == "faq"), None)
    sss_basligi = None
    if i_sss is not None and i_sss > 0 and v["akis"][i_sss - 1][0] == "h3":
        sss_basligi = v["akis"][i_sss - 1][1]
    atla = {i for i, (t, _) in enumerate(v["akis"]) if t == "fix"}
    if i_sss is not None:
        atla.add(i_sss)
        if sss_basligi:
            atla.add(i_sss - 1)

    for i, (t, d) in enumerate(v["akis"]):
        if i in atla:
            continue
        if t in ("h2", "h3", "h4"):
            d = BASLIK_REVIZE.get((slug, re.sub(r"<[^>]+>", "", d).strip()), d)
            parcalar.append(f"<{t}>{d}</{t}>"); kaynak.append(d)
            if t == "h2":
                _kt = KATEGORI_TANIM.get(slug)
                if _kt:
                    parcalar.append(f"<p><strong>{_kt[0]}</strong> {_kt[1]}</p>")
            if t == "h4":
                # Alt turun hemen altina "tanim-once" cumlesi: AI Overview ve "X nedir"
                # sorgulari bu kalibi dogrudan cekiyor.
                _ad = re.sub(r"<[^>]+>", "", d).strip()
                _tanim = ALT_TUR_TANIM.get((slug, _ad))
                if _tanim:
                    parcalar.append(f"<p><strong>{_ad} nedir?</strong> {_tanim}</p>")
        elif t == "p":
            parcalar.append(f"<p>{d}</p>"); kaynak.append(d)
        elif t == "ul":
            parcalar.append(render_list(d)); kaynak.extend(d)
        elif t == "ol":
            parcalar.append("<ol>\n" + "\n".join(f"  <li>{x}</li>" for x in d) + "\n</ol>")
            kaynak.extend(d)
        elif t == "tldr":
            parcalar.append(render_tldr(v["tldr"]))
        elif t == "info":
            parcalar.append(render_info_card(v["info"]))
        elif t == "tablo":
            b, sat = v["tablolar"][d]
            parcalar.append(render_table(b, sat))
        elif t == "not":
            parcalar.append(render_editor_note(v["notlar"][d]))
        elif t == "callout":
            parcalar.append(render_highlight(v["callout"]))
        elif t == "dyn" and v["dyn"]:
            b, a = v["dyn"]
            parcalar.append(render_category_dynamic_cta(b, a, dyn_url))

    # Zorunlu bilesen yedekleri
    if not v["callout"]:
        # lisans hatirlatmasi kapanistan once, Fix CTA'nin hemen ustune
        parcalar.append(render_highlight(LISANS_METNI))
    if not v["dyn"]:
        b, a = DYN_CTA_METNI.get(slug, _VARSAYILAN_DYN)
        # sayfa ortasina: son ucte birin basina yerlestir
        yer = max(1, int(len(parcalar) * 0.55))
        parcalar.insert(yer, render_category_dynamic_cta(b, a, dyn_url))

    if v["fix"]:
        parcalar.append(render_category_fix_cta(*v["fix"]))
    if v["faq"]:
        faq = [(FAQ_SORU_REVIZE.get((slug, q), q), a) for q, a in v["faq"]]
        faq += EK_FAQ.get(slug, [])          # alt tur ve arama niyeti odakli yeni sorular
        _kaldir = FAQ_KALDIR.get(slug, [])
        faq = [(q, a) for q, a in faq if q not in _kaldir]
        v["faq_revize"] = [(q, y) for (q, _), (y, _) in zip(v["faq"], faq) if q != y]
        parcalar.append(f"<h3>{sss_basligi or 'Sık Sorulan Sorular'}</h3>")
        parcalar.append(render_faq_accordion(faq))
        parcalar.append(render_faq_schema(faq))   # schema GORUNEN soruyla birebir
    _h2 = next((d for t, d in v["akis"] if t == "h2"), slug)
    _ad = re.sub(r"<[^>]+>", "", _h2).strip()
    _kt = KATEGORI_TANIM.get(slug)
    parcalar.append(render_kategori_schema(
        kategori_url(slug), _ad, _kt[1] if _kt else _ad, guncelleme=URETIM_TARIHI))
    return "\n".join(parcalar), kaynak


def docx_yaz(html, yol):
    """Doc SADECE HTML icerir (kullanici tercihi); blog doc standardiyla ayni bicim."""
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11); st.font.color.rgb = RGBColor(0x10, 0x33, 0x2F)
    for satir in html.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(satir)
        r.font.name = "Courier New"; r.font.size = Pt(7)
    doc.save(yol)



# --- Oyun sayisi yuvarlama (skill kurali + markanin Sheet5 notu) ---
# ">=100 icin en yakin 100 asagi, <100 icin en yakin 10 asagi". Zaten "+" tasiyan degere DOKUNULMAZ.
_SAYAC = (r"(oyun|oyunu|oyunla|oyuna|oyunun|oyunlar|oyunları|yapım|yapımı|yapıma|yapımın|"
          r"yapımlar|yapımları|demo|demoyu|başlık|indie|İndie|MOBA|adet)")


def _yuvarla(n):
    n = int(n)
    return f"{(n // 100) * 100}+" if n >= 100 else f"{max((n // 10) * 10, 10)}+"


def sayilari_yuvarla(html):
    """info-card 'Kütüphane Boyutu' degerini ve govdede ayni sayiyla kurulan 'N oyun' kaliplarini
    yuvarlar. Donus: (html, degisiklikler)."""
    govde_bas = html.find("</style>") + len("</style>")
    bas, govde = html[:govde_bas], html[govde_bas:]
    m = re.search(r'<div class="gp-cell-value">(\d+)(?![\d.+])([^<]*)</div>\s*'
                  r'<div class="gp-cell-label">Kütüphane Boyutu</div>', govde)
    if not m:
        return html, []
    ham, yeni = m.group(1), _yuvarla(m.group(1))
    degisim = [(f"{ham}{m.group(2)}", f"{yeni}{m.group(2)}", "info-card")]
    govde = govde[:m.start(1)] + yeni + govde[m.end(1):]
    # Sayi ile sayilan sozcuk arasinda sifat olabilir: "110 aile dostu oyun", "14 MOBA oyunu"
    desen = re.compile(r"(?<![\d.,+])" + ham + r"(?![\d.+])"
                       r"(\s+(?:[A-Za-zÇĞİÖŞÜçğıöşü+.-]+\s+){0,3})" + _SAYAC)
    for mm in desen.finditer(govde):
        degisim.append((f"{ham}{mm.group(1)}{mm.group(2)}", f"{yeni}{mm.group(1)}{mm.group(2)}", "gövde"))
    govde = desen.sub(yeni + r"\1\2", govde)
    return bas + govde, degisim


# --- Kaynaktan tasinan dizgi hatalari (nokta sonrasi bosluk, bozuk cumle) ---
def dizgi_duzelt(html):
    govde_bas = html.find("</style>") + len("</style>")
    bas, govde = html[:govde_bas], html[govde_bas:]
    duzeltme = []
    # "alınır.) doğrudan çalışır." -> kaynaktan gelen bozuk cumle
    if "alınır.) doğrudan çalışır." in govde:
        govde = govde.replace("alınır.) doğrudan çalışır.", "alınır.")
        duzeltme.append("bozuk cümle: 'alınır.) doğrudan çalışır.' -> 'alınır.'")
    # cumle sonu noktadan sonra bosluk yok: "desteklenir.Skyrim" -> "desteklenir. Skyrim"
    yeni, n = re.subn(r"([a-zçğıöşü])\.([A-ZÇĞİÖŞÜ][a-zçğıöşü])", r"\1. \2", govde)
    if n:
        govde = yeni
        duzeltme.append(f"nokta sonrası boşluk: {n} yer")
    # "Öne Çıkan Yapımlar:</strong><em>Cyberpunk" -> iki etiket arasinda bosluk yok, ekranda
    # "Yapımlar:Cyberpunk" gibi yapisik goruluyor.
    yeni, n2 = re.subn(r"([:.])</(strong|b|em|i)><(strong|b|em|i|span)(\s|>)",
                       r"\1</\2> <\3\4", govde)
    if n2:
        govde = yeni
        duzeltme.append(f"etiketler arası boşluk: {n2} yer")
    return bas + govde, duzeltme


def isle(slug):
    ham = open(os.path.join(KAYNAK, f"{slug}.html"), encoding="utf-8").read()
    v = cikar(BeautifulSoup(ham, "html.parser"))
    body, kaynak = kur(v, slug)
    # Kural 20: govde ici kategori linki. Kategori KENDINE link vermez.
    # Firsatlar yonlendirmesi kaldirildi: yazarin cumlesi KALIR, yalniz <a> sarmalayicisi cozulur.
    firsat_sayisi = len(re.findall(r'<a[^>]*firsatlar[^>]*>', body))
    body = re.sub(r'<a[^>]*href="[^"]*firsatlar[^"]*"[^>]*>(.*?)</a>', r'\1', body, flags=re.S)
    mevcut = set(re.findall(r'href="(https://gameplus\.com\.tr/gfn/oyunlar/[a-z-]+)"', body))
    body, kat_linkler = auto_link_categories(
        body, max_links=2, haric=tuple(mevcut | {kategori_url(slug)}))
    body, mod_degisim = mod_iddialarini_kaldir(body, slug)
    body, link_sayac = apply_link_policy(body)          # Kural 21: dis nofollow, hepsi yeni sekme
    final = wrap_gp_content(CATEGORY_STYLE + "\n" + group_into_sections(body))  # Kural 22
    final, yazim = yazim_normalize(final)
    final, sayi_degisim = sayilari_yuvarla(final)
    final, dizgi = dizgi_duzelt(final)

    oyun_sayisi = next((dg for _, dg in v["info"] if "+" in dg), None)
    sonuc = verify_category_output(final, oyun_sayisi_metni=oyun_sayisi)
    hatalar = [x for x in sonuc if x[0] == "FAIL"]
    orij = "".join(f"<p>{x}</p>" for x in kaynak)
    orij, _ = mod_iddialarini_kaldir(orij, slug)
    orij, _ = yazim_normalize(orij)
    ok_kaynak, eksik, oran = verify_source_preserved(orij, final)

    os.makedirs(CIKTI, exist_ok=True)
    if not hatalar and ok_kaynak:
        docx_yaz(final, os.path.join(CIKTI, f"gfn-oyunlar-{slug}.docx"))
        open(os.path.join(CIKTI, f"gfn-oyunlar-{slug}.txt"), "w", encoding="utf-8").write(final)
    return {
        "slug": slug, "ok": not hatalar and ok_kaynak,
        "hatalar": [f"{x[1]}: {x[2]}" for x in hatalar],
        "kaynak_oran": oran, "eksik": eksik[:3],
        "boyut": len(final), "tldr": len(v["tldr"]), "info": len(v["info"]),
        "tablo": len(v["tablolar"]), "faq": len(v["faq"]),
        "cta": CTA_HEDEF.get(slug, PAKETLER).rsplit("/", 1)[-1],
        "linkler": [i for i, _ in kat_linkler],
        "oyun_sayisi": oyun_sayisi, "firsatlar_cozuldu": firsat_sayisi, "faq_revize": v.get("faq_revize", []), "sayi": sayi_degisim, "dizgi": dizgi, "mod": mod_degisim, "yazim": yazim, "link": link_sayac,
        "fix_baslik": (v["fix"] or ("", ""))[0], "dyn_baslik": (v["dyn"] or ("", ""))[0],
    }


if __name__ == "__main__":
    hedef = sys.argv[1:] or sorted(x[:-5] for x in os.listdir(KAYNAK) if x.endswith(".html"))
    basarili, basarisiz = [], []
    for s in hedef:
        r = isle(s)
        (basarili if r["ok"] else basarisiz).append(r)
        durum = "OK " if r["ok"] else "FAIL"
        print(f"{durum} {r['slug']:20s} {r['boyut']:6d} char | tldr {r['tldr']} info {r['info']} "
              f"tablo {r['tablo']} faq {r['faq']} | cta {r['cta']:18s} | metin %{r['kaynak_oran']*100:.0f} "
              f"| link {r['linkler'] or '-'}")
        for h in r["hatalar"]:
            print(f"      ! {h}")
        if r["eksik"]:
            print(f"      ! eksik metin: {r['eksik']}")
    print(f"\n{len(basarili)}/{len(hedef)} başarılı | {len(basarisiz)} hatalı")
