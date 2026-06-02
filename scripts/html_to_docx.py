"""Convert simulasyon-icerik.html to a properly styled .docx file."""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from html.parser import HTMLParser

INPUT_HTML = "/Users/Erdo/Desktop/Claude Projects/Dispatch/simulasyon-icerik.html"
OUTPUT_DOCX = "/Users/Erdo/Desktop/Claude Projects/Game+ /İçerikler/gameplus.com.tr-gfn-oyunlar-simulasyon.docx"

class HTMLToDocx(HTMLParser):
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.tag_stack = []
        self.current_para = None
        self.current_list_type = None
        self.in_table = False
        self.current_table = None
        self.current_row = None
        self.current_cell_para = None
        self.is_header_row = False
        self.in_link = False
        self.link_url = None
        self.in_bold = False
        self.text_buffer = []
        self.collected_for_link = []

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        self.tag_stack.append(tag)

        if tag in ("h2", "h3", "h4"):
            level = int(tag[1])
            self.current_para = self.doc.add_heading("", level=level)
        elif tag == "p":
            if self.in_table:
                pass  # Skip <p> inside <td>
            else:
                self.current_para = self.doc.add_paragraph()
        elif tag == "ul" or tag == "ol":
            self.current_list_type = tag
        elif tag == "li":
            style = "List Number" if self.current_list_type == "ol" else "List Bullet"
            self.current_para = self.doc.add_paragraph(style=style)
        elif tag == "table":
            self.in_table = True
            self.current_table = self.doc.add_table(rows=0, cols=0)
            self.current_table.style = "Light Grid Accent 1"
        elif tag == "thead":
            self.is_header_row = True
        elif tag == "tbody":
            self.is_header_row = False
        elif tag == "tr":
            if self.in_table:
                self.current_row = self.current_table.add_row() if self.current_table.rows else None
                if self.current_row is None:
                    # First row, add columns later
                    pass
        elif tag in ("th", "td"):
            # Will be handled in handle_data
            pass
        elif tag == "a":
            self.in_link = True
            self.link_url = attrs_dict.get("href", "")
        elif tag == "strong" or tag == "b":
            self.in_bold = True

    def handle_endtag(self, tag):
        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()

        if tag == "a":
            self.in_link = False
            self.link_url = None
        elif tag == "strong" or tag == "b":
            self.in_bold = False
        elif tag == "ul" or tag == "ol":
            self.current_list_type = None
        elif tag == "table":
            self.in_table = False
        elif tag == "tr":
            self.current_row = None
        elif tag in ("p", "li", "h2", "h3", "h4"):
            self.current_para = None

    def handle_data(self, data):
        text = data.strip("\n")
        if not text.strip() and not self.current_para:
            return

        # Handle table cells
        if self.in_table and self.tag_stack:
            inside_th = "th" in self.tag_stack
            inside_td = "td" in self.tag_stack
            if inside_th or inside_td:
                tbl = self.current_table
                # Build columns on the fly
                if self.is_header_row and inside_th:
                    if tbl.rows and len(tbl.rows) == 1:
                        row = tbl.rows[0]
                    else:
                        # Create header row if needed
                        if not tbl.rows:
                            tbl.add_row()
                        row = tbl.rows[0]
                    # Add a new cell column
                    cell_idx = len([c for c in row.cells if c.text]) if False else None
                    # Simpler: track via columns count
                    if not tbl.columns or len(row.cells) <= sum(1 for c in row.cells if c.text):
                        # Add column
                        if len(tbl.columns) == 0:
                            tbl.add_column(width=Pt(100))
                    # Easier: just append to next empty cell
                    for c in row.cells:
                        if not c.text:
                            run = c.paragraphs[0].add_run(text)
                            run.bold = True
                            return
                    # If all filled, add new column
                    tbl.add_column(width=Pt(100))
                    new_cell = row.cells[-1]
                    run = new_cell.paragraphs[0].add_run(text)
                    run.bold = True
                    return
                elif not self.is_header_row and inside_td:
                    # body row: find or create row, then fill next empty cell
                    if not tbl.rows or self.current_row is None:
                        # Need to start a new row
                        if tbl.columns and len(tbl.columns) > 0:
                            new_row = tbl.add_row()
                            self.current_row = new_row
                        else:
                            return
                    # Find next empty cell
                    for c in self.current_row.cells:
                        if not c.text:
                            run = c.paragraphs[0].add_run(text)
                            if self.in_bold:
                                run.bold = True
                            return
                    return

        # Normal paragraph
        if self.current_para is None and text.strip():
            self.current_para = self.doc.add_paragraph()

        if self.current_para is not None and text.strip():
            run = self.current_para.add_run(text)
            if self.in_link:
                run.font.color.rgb = RGBColor(0x06, 0x5F, 0xD8)
                run.font.underline = True
            if self.in_bold:
                run.bold = True


# Simpler approach: use BeautifulSoup-like parsing
from bs4 import BeautifulSoup

def html_to_docx(html_path, docx_path):
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    # Add title from H2
    h2 = soup.find("h2")
    if h2:
        title = doc.add_heading(h2.get_text(), level=1)

    # Process each top-level element in order
    body_elements = soup.find_all(["h2", "h3", "h4", "p", "ul", "ol", "table"], recursive=False)
    # If no recursive=False matches (because they're nested), grab all
    if not body_elements:
        body_elements = list(soup.children)

    seen_first_h2 = False
    for elem in soup.children:
        if not hasattr(elem, "name") or elem.name is None:
            continue
        name = elem.name
        if name == "h2":
            if seen_first_h2:
                doc.add_heading(elem.get_text(), level=2)
            else:
                seen_first_h2 = True  # Already added as title
        elif name == "h3":
            doc.add_heading(elem.get_text(), level=2)
        elif name == "h4":
            doc.add_heading(elem.get_text(), level=3)
        elif name == "p":
            p = doc.add_paragraph()
            add_inline(p, elem)
        elif name == "ul":
            for li in elem.find_all("li", recursive=False):
                bul = doc.add_paragraph(style="List Bullet")
                add_inline(bul, li)
        elif name == "ol":
            for li in elem.find_all("li", recursive=False):
                num = doc.add_paragraph(style="List Number")
                add_inline(num, li)
        elif name == "table":
            add_table(doc, elem)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def add_inline(paragraph, element):
    """Walk inline children of a block element and add runs preserving formatting."""
    for child in element.children:
        if isinstance(child, str):
            txt = child.replace("\n", " ").rstrip()
            if txt:
                paragraph.add_run(txt)
        elif hasattr(child, "name"):
            if child.name in ("strong", "b"):
                # Recurse: bold runs
                add_inline_with_format(paragraph, child, bold=True)
            elif child.name == "a":
                href = child.get("href", "")
                # Add link as colored underlined text + URL in brackets is too noisy
                # Just style the anchor text
                add_inline_with_format(paragraph, child, link=True, href=href)
            elif child.name == "em" or child.name == "i":
                add_inline_with_format(paragraph, child, italic=True)
            elif child.name == "br":
                paragraph.add_run("\n")
            else:
                add_inline(paragraph, child)


def add_inline_with_format(paragraph, element, bold=False, italic=False, link=False, href=""):
    """Add element children to paragraph with formatting."""
    for child in element.children:
        if isinstance(child, str):
            txt = child.replace("\n", " ").rstrip()
            if txt:
                run = paragraph.add_run(txt)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if link:
                    run.font.color.rgb = RGBColor(0x06, 0x5F, 0xD8)
                    run.font.underline = True
        elif hasattr(child, "name"):
            if child.name in ("strong", "b"):
                add_inline_with_format(paragraph, child, bold=True or bold, italic=italic, link=link, href=href)
            elif child.name == "a":
                inner_href = child.get("href", href)
                add_inline_with_format(paragraph, child, bold=bold, italic=italic, link=True, href=inner_href)
            elif child.name == "em" or child.name == "i":
                add_inline_with_format(paragraph, child, bold=bold, italic=True, link=link, href=href)
            else:
                add_inline_with_format(paragraph, child, bold=bold, italic=italic, link=link, href=href)


def add_table(doc, table_elem):
    # Count columns from first row
    rows = table_elem.find_all("tr")
    if not rows:
        return
    first_row_cells = rows[0].find_all(["th", "td"])
    num_cols = len(first_row_cells)
    if num_cols == 0:
        return

    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.style = "Light Grid Accent 1"

    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for c_idx, cell in enumerate(cells):
            if c_idx >= num_cols:
                break
            table_cell = tbl.cell(r_idx, c_idx)
            # Clear existing paragraph and add our content
            table_cell.text = ""
            p = table_cell.paragraphs[0]
            is_header = cell.name == "th"
            add_inline_with_format(p, cell, bold=is_header)


if __name__ == "__main__":
    html_to_docx(INPUT_HTML, OUTPUT_DOCX)
